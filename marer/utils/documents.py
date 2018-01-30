from copy import deepcopy
from io import BytesIO
from string import Formatter
import re

from django.core.files.base import ContentFile
from docx import Document
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from openpyxl.writer.excel import save_virtual_workbook
from lxml import etree

from marer.models import Issue


def fill_xlsx_file_with_issue_data(path: str, issue: Issue) -> ContentFile:
    wb = load_workbook(path, keep_vba=True, guess_types=False)
    self_data = issue.__dict__
    self_data['user'] = issue.user
    self_data['issue'] = issue
    for ws_name in wb.sheetnames:
        ws = wb.get_sheet_by_name(ws_name)
        for cell in ws.get_cell_collection():
            if cell.value is not None:
                old_cell_value = str(cell.value)
                try:
                    new_cell_value = old_cell_value.format(**self_data)
                except KeyError:
                    new_cell_value = old_cell_value
                if new_cell_value != old_cell_value:
                    cell.value = new_cell_value
                cell.value = new_cell_value
    cf = ContentFile(save_virtual_workbook(wb))
    wb.close()
    return cf


def _optimize_paragraph(paragraph: Paragraph):
    runs = paragraph.runs
    paragraph_params = Formatter().parse(paragraph.text)
    paragraph_params = [pp for pp in paragraph_params if pp[1] is not None]
    if len(paragraph_params) == 0:
        return

    for run in runs[:-1]:
        curr_idx = runs.index(run)
        next_run = runs[curr_idx + 1]
        open_brace_last_pos = str(run.text).rfind('{')
        close_brace_last_pos = str(run.text).rfind('}')

        if open_brace_last_pos > close_brace_last_pos:
            next_run.text = run.text[open_brace_last_pos:] + next_run.text
            run.text = run.text[:open_brace_last_pos]


def _fill_docx_paragraph_with_dict(paragraph: Paragraph, data: dict) -> None:
    runs = paragraph.runs
    _optimize_paragraph(paragraph)

    for run in runs:
        old_text = str(run.text)
        try:
            if '{' in old_text:
                for found in set(re.findall('{[^}]*?}', old_text)):
                    old_text = old_text.replace(found, '^%s^' % found)
                new_text = old_text.format(**data)
                new_text = new_text.replace('^^', 'нет').replace('^', '')
            else:
                new_text = old_text
        except KeyError:
            pass
        except ValueError:
            pass
        except IndexError:
            pass
        else:
            if old_text != new_text:
                run.text = new_text


def fill_docx_file_with_issue_data(path: str, issue: Issue) -> ContentFile:
    self_data = issue.__dict__
    self_data['user'] = issue.user
    self_data['issue'] = issue

    doc = Document(path)
    helper = WordDocumentHelper()
    doc = helper.prepare(doc, self_data)
    for paragraph in doc.paragraphs:
        _fill_docx_paragraph_with_dict(paragraph, self_data)
    for table in doc.tables:
        rows_cnt = len(table.rows)
        for row_idx in range(0, rows_cnt):
            for cell in table.row_cells(row_idx):
                for p in cell.paragraphs:
                    _fill_docx_paragraph_with_dict(p, self_data)

    stream = BytesIO()
    doc.save(stream)

    stream.seek(0)
    cf = ContentFile(stream.read())
    return cf


class WordDocumentHelper:
    TR_OFFSET = 2
    TD_OFFSET = 1

    def get_cells(self, table, row_i):
        """ Возвращает только элементы-ячейки (первый элемент - не нужный)  """
        data = table._tbl[row_i][self.TD_OFFSET:]
        return data

    def get_text(self, item):
        return ''.join(item.xpath(".//*/text()"))

    def get_count(self, data, field):
        value = data
        for path in field.split('.'):
            if isinstance(value, dict):
                value = value.get(path)
            elif hasattr(value, path):
                if callable(getattr(value, path)):
                    value = getattr(value, path)()
                else:
                    value = getattr(value, path)
        return len(value or [])

    def prepare(self, doc, data):
        """ Обработка циклов для word документа
        :param doc: редактируемый документ
        :param data: данные для вставки, нужны для расчета количества элементов в списках
        :return:
        """
        for table in doc.tables:
            # сначала все оптимизируем, иначе потом это делать трудоемко
            for row_idx in range(len(table.rows)):
                for c in table.row_cells(row_idx + self.TR_OFFSET):
                    for p in c.paragraphs:
                        _optimize_paragraph(p)
            # ищем признаки начала и конца цикла
            data_for_replace = []
            rows_cnt = len(table._tbl)
            row_idx = self.TR_OFFSET
            while row_idx < rows_cnt:
                end_row = None  # конечная строка для копирования
                # ищем в первой ячейке признак начала цикла
                first_cell = self.get_cells(table, row_idx)[0]
                found = re.findall(r'{[^{}]*\|for}', self.get_text(first_cell))
                if found:  # как только нашли начало цикла, ищем его конец
                    found = found[0][1:-5]
                    start_row = row_idx
                    for row_i in range(row_idx, rows_cnt):
                        for cell in self.get_cells(table, row_i):
                            pattern = r'{%s*\|endfor}' % found
                            found_end = re.findall(pattern, self.get_text(cell))
                            if found_end:
                                end_row = row_i
                                break
                        if end_row:
                            break
                    if not end_row:
                        end_row = start_row
                    # нашли конец, теперь клонируем строки с start_end по end_row включительно
                    position = end_row
                    cloned_rows = end_row - start_row
                    count = self.get_count(data, found)
                    for i in range(count - 1):
                        for j in range(start_row, end_row + 1):
                            clone = deepcopy(table._tbl[j])
                            table._tbl[position].addnext(clone)
                            rows_cnt += 1
                            position += 1
                            cloned_rows += 1
                    # сохраняем данные для будущей замены
                    data_for_replace.append([
                        start_row,
                        end_row - start_row + 1,
                        found,
                        cloned_rows + 1,  # количество строк, в которых возможно нужна замена obj
                        count,  # количество элементов массива
                    ])
                    row_idx = position + 1

                row_idx += 1
            # собственно замена
            for data in data_for_replace:
                replace_dict = {
                    "{%s|for}" % data[2]: '',
                    "{%s|endfor}" % data[2]: ''
                }
                # пробегамся только по отмеченным участкам для замены
                for row_idx in range(data[0], data[0] + data[3]):
                    index = (row_idx - data[0]) // data[1]  # индекс для вставки
                    for i, cell in enumerate(self.get_cells(table, row_idx)):
                        cell_text = etree.tostring(cell)
                        old_text = cell_text
                        new_text = old_text
                        # скрываем признаки начала и конца цикла, заменяем obj на переменную с индексом
                        for key, value in replace_dict.items():
                            new_text = new_text.replace(bytes(key, encoding='utf-8'), bytes(value, encoding='utf-8'))
                        if data[4]: # если данные имеются,
                            new_text = new_text.replace(bytes('{obj', encoding='utf-8'),
                                                        bytes('{' + data[2] + '[%s]' % index, encoding='utf-8'))
                        else:
                            found = re.findall('{obj[^{}]*}', str(old_text))
                            for f in found:
                                new_text = new_text.replace(bytes(f, encoding='utf-8'),
                                                            bytes('нет', encoding='utf-8'))
                        # заменяем содержимое ячейки на новое содержимое
                        table._tbl[row_idx][i + 1] = etree.fromstring(new_text)

        return doc